"""Weekly report parser - extracts content from docx with charts preserved."""
import os
import base64
import io
from pathlib import Path
from docx import Document

# Use Obsidian path locally, ../data/weekly on cloud
WEEKLY_DIR = Path(os.getenv(
    "WEEKLY_REPORTS_DIR",
    "/Users/rhea/Obsidian data/work/周报" if not (os.getenv("HF_SPACE") or os.getenv("RENDER")) else "../data/weekly"
))


def list_weekly_reports():
    """List all available weekly reports."""
    reports = []
    if not WEEKLY_DIR.exists():
        return reports
    
    for d in sorted(WEEKLY_DIR.iterdir(), reverse=True):
        if d.is_dir() and d.name.endswith("期"):
            docx_files = list(d.glob("*.docx"))
            for df in docx_files:
                reports.append({
                    "period": d.name,
                    "filename": df.name,
                    "path": str(df),
                })
    return reports


def extract_weekly_report(filepath: str):
    """Extract content from a weekly report docx, preserving images as base64."""
    doc = Document(filepath)
    
    # Extract images
    images = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                ext = rel.target_part.partname.split(".")[-1]
                b64 = base64.b64encode(img_data).decode()
                images[rel_id] = f"data:image/{ext};base64,{b64}"
            except Exception:
                pass
    
    # Build HTML content preserving structure
    html_parts = []
    img_idx = 0
    image_keys = list(images.keys())
    
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        
        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            
            if para is None:
                continue
            
            text = para.text.strip()
            if not text and not _has_image(para):
                continue
            
            # Check for inline images
            drawing_elements = element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
            drawing_elements += element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
            
            for drawing in drawing_elements:
                blip = drawing.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                for b in blip:
                    embed = b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed and embed in images:
                        html_parts.append(f'<div style="text-align:center;margin:16px 0"><img src="{images[embed]}" style="max-width:100%;border:1px solid #eee;border-radius:4px"></div>')
            
            # Style based on content
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "标题" in style_name:
                level = 2
                for c in style_name:
                    if c.isdigit():
                        level = int(c)
                        break
                html_parts.append(f"<h{level}>{text}</h{level}>")
            elif text:
                # Check if it's a section header
                if text.startswith("【") and text.endswith("】"):
                    html_parts.append(f'<h2 style="color:#409eff;margin-top:24px;border-bottom:2px solid #409eff;padding-bottom:8px">{text}</h2>')
                elif any(kw in text for kw in ["产业链价格走势", "价格分析", "供需及市场分析"]):
                    html_parts.append(f'<h3 style="margin-top:16px;color:#303133">{text}</h3>')
                else:
                    html_parts.append(f"<p>{text}</p>")
        
        elif tag == "tbl":
            # Table
            table_idx = None
            for i, t in enumerate(doc.tables):
                if t._element is element:
                    table_idx = i
                    break
            
            if table_idx is not None:
                table = doc.tables[table_idx]
                html_parts.append('<div style="overflow-x:auto;margin:12px 0"><table style="border-collapse:collapse;width:100%;font-size:13px">')
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        is_header = any(run.bold for run in cell.paragraphs[0].runs if run.text) if cell.paragraphs[0].runs else False
                        tag_type = "th" if (is_header or row is table.rows[0]) else "td"
                        style = "border:1px solid #e4e7ed;padding:6px 10px;text-align:left;background:#f5f7fa" if tag_type == "th" else "border:1px solid #e4e7ed;padding:6px 10px"
                        html_parts.append(f'<{tag_type} style="{style}">{cell.text}</{tag_type}>')
                    html_parts.append("</tr>")
                html_parts.append("</table></div>")
    
    content = "\n".join(html_parts) if html_parts else "<p>暂无内容</p>"
    return content


def _has_image(para) -> bool:
    """Check if paragraph contains images."""
    drawings = para._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
    drawings += para._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
    return len(drawings) > 0
